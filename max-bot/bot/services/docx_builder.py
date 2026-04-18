"""Генерация Word-документа (.docx) с транскрибацией."""

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def build_docx(
    text: str,
    output_path: Path,
    duration: str = "",
    original_filename: str = "",
    language: str = "ru",
    correction_applied: bool = False,
) -> Path:
    """Создать Word-документ с транскрибацией."""
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Стиль по умолчанию
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Заголовок
    title = doc.add_heading("ТРАНСКРИБАЦИЯ ИНТЕРВЬЮ", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
        run.font.name = "Arial"

    # Имя файла
    if original_filename:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(original_filename)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Метаданные
    date_str = datetime.now().strftime("%d.%m.%Y %H:%M")
    meta_lines = [
        f"Дата: {date_str}",
        f"Длительность: {duration}",
        f"STT: Wejet ИИ сервисы",
    ]
    if correction_applied:
        meta_lines.append("Коррекция: Wejet ИИ сервисы ✓")

    for line in meta_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Разделитель
    p = doc.add_paragraph()
    run = p.add_run("─" * 70)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)

    # Основной текст
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        p = doc.add_paragraph()

        # Формат "Спикер N: текст" или "Участник N: текст"
        speaker_prefix = ""
        remaining = line
        for i in range(1, 20):
            for tmpl in (f"Спикер {i}:", f"Участник {i}:"):
                if line.startswith(tmpl):
                    speaker_prefix = tmpl
                    remaining = line[len(tmpl):].strip()
                    break
            if speaker_prefix:
                break

        if speaker_prefix:
            run = p.add_run(speaker_prefix + " ")
            run.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
            run.font.size = Pt(12)
            run.font.name = "Arial"

        # Временная метка [ЧЧ:ММ:СС]
        timestamp = ""
        text_after = remaining
        if remaining.startswith("[") and "]" in remaining:
            bracket_end = remaining.index("]")
            possible_ts = remaining[1:bracket_end]
            if ":" in possible_ts and all(c.isdigit() or c == ":" for c in possible_ts):
                timestamp = f"[{possible_ts}] "
                text_after = remaining[bracket_end + 1:].strip()

        if timestamp:
            run = p.add_run(timestamp)
            run.font.color.rgb = RGBColor(0x96, 0x96, 0x96)
            run.font.size = Pt(9)

        if text_after:
            run = p.add_run(text_after)
            run.font.size = Pt(12)
            run.font.name = "Arial"
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Футер
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("─" * 70)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)

    footer = "Транскрибация: Wejet ИИ сервисы."
    if correction_applied:
        footer += " Коррекция: Wejet ИИ сервисы."
    footer += " Возможны неточности."

    p = doc.add_paragraph()
    run = p.add_run(footer)
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(str(output_path))
    logger.info(f"DOCX создан: {output_path}")
    return output_path
